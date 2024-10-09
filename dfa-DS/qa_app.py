import streamlit as st
from langchain_google_genai import GoogleGenerativeAI
from langchain.output_parsers import PydanticOutputParser
from langchain.chains import LLMChain
from langchain.schema import Document
from langchain.prompts import PromptTemplate
from PyPDF2 import PdfReader
import docx
from pydantic import BaseModel, Field
from typing import List
import pandas as pd
from io import StringIO

st.set_page_config(page_title="Teacher question and answer practice", page_icon="✍")
# Define the Pydantic Model
class QAParser(BaseModel):
    questions: List[str] = Field(..., description="List of questions generated")
    answers: List[str] = Field(..., description="List of answers corresponding to each question")


# Define the Prompt Template
prompt_template_str = """
Please generate {number} questions and their corresponding answers based on the following context:

Context:
{context}

Please provide the output in JSON format following this structure:

{format_instructions}
"""

prompt_template = PromptTemplate(
    template=prompt_template_str,
    input_variables=["number", "context", "format_instructions"]
)


# Define the Output Parser
parser = PydanticOutputParser(pydantic_object=QAParser)


# Initialize the LLM
# Using GoogleGenerativeAI with model "gemini-pro"
llm = GoogleGenerativeAI(model="gemini-pro")


# Initialize the LLMChain
chain = LLMChain(
    llm=llm,
    prompt=prompt_template,
    output_parser=parser
)


# Custom Document Loader
class LambdaStreamlitLoader:
    def __init__(self, file) -> None:
        self.file = file

    def lazy_load(self):
        file_name = self.file.name
        *_, ext = file_name.split(".")
        ext = ext.lower()

        if ext in ["docx", "doc"]:
            doc = docx.Document(self.file)
            for paragraph in doc.paragraphs:
                lines = paragraph.text.split("\n")
                for line in lines:
                    line = line.strip()
                    if line:
                        yield Document(page_content=line)

        elif ext == "pdf":
            reader = PdfReader(self.file)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    lines = text.split("\n")
                    for line in lines:
                        line = line.strip()
                        if line:
                            yield Document(page_content=line)
        else:
            st.error("Unsupported file format!")

# function to create csv from qa_data
def create_csv(qa_data):
    qa_df = pd.DataFrame({
        "Question": qa_data.questions,
        "Answer": qa_data.answers
    })
    
    csv_buffer = StringIO()
    qa_df.to_csv(csv_buffer, index=False)

    return csv_buffer.getvalue()

# Template for guiding the student on their answers
help_template = """
I have been asked this question: {question}
And then I provided this answer: {student_answer}
However, this is the actual answer: {ground_truth}
You are not to respond with the actual answer, just compare my answer to the right answer and guide me.
If my answer indicates that I don't know, guide me with a clearer step-by-step approach to make me understand it, with words of encouragement.
The answer doesn't have to be verbatim, tell me I am correct if I use other words that have the same meaning.
"""

# Initialize LLM for help evaluation
help_llm = GoogleGenerativeAI(model="gemini-pro")


# Initialize session state variables
if "qa_data" not in st.session_state:
    st.session_state.qa_data = None

if "count" not in st.session_state:
    st.session_state.count = 0

if "questions_generated" not in st.session_state:
    st.session_state.questions_generated = False

if "document" not in st.session_state:
    st.session_state.document = None


# Sidebar for File Upload and Question Generation
with st.sidebar:
    st.header("Upload and Generate Q&A")
    file = st.file_uploader("Upload document here", type=["pdf", "docx", "doc"])

    if file:
        if st.session_state.document is None:
            loader = LambdaStreamlitLoader(file)
            st.session_state.document = list(loader.lazy_load())
            st.success("Document Loaded successfully")

    if file and st.session_state.document:
        number_ip = st.number_input(
            "How many questions would you like to set?",
            step=1,
            min_value=1,
            format="%i",
            value=5
        )

        if st.button("Generate") and st.session_state.document:
            with st.spinner("Generating questions..."):
                context = "\n".join(doc.page_content for doc in st.session_state.document)
                try:
                    qa_result = chain.run(
                        number=number_ip,
                        context=context,
                        format_instructions=parser.get_format_instructions()
                    )
                    st.session_state.qa_data = qa_result
                    st.session_state.questions_generated = True
                    st.session_state.count = 0
                    st.success("Questions and answers generated successfully!")
                except Exception as e:
                    st.error(f"An error occurred during generation: {e}")

# Main Content Area
st.title("Question and Answer Session")

if st.session_state.questions_generated and st.session_state.qa_data:
    qa_data = st.session_state.qa_data
    if qa_data:
        with st.sidebar:
            csv_data = create_csv(st.session_state.qa_data)

            st.download_button(
                label="Download Q&A as CSV",
                data=csv_data,
                file_name="qa_data.csv",
                mime="text/csv"
            )
            
    questions = qa_data.questions
    answers = qa_data.answers
    

    if st.session_state.count < len(questions):
        current_index = st.session_state.count
        current_question = questions[current_index]
        current_answer = answers[current_index]

        st.markdown(f"### Question {current_index + 1}/{len(questions)}")
        st.write(f"**Question:** {current_question}")

        with st.expander("Reveal Answer"):
            st.write("**Actual answer from material:**")
            st.write(current_answer)

        # Initialize user answer in session state
        if f"user_answer_{current_index}" not in st.session_state:
            st.session_state[f"user_answer_{current_index}"] = ""

        # Text area for user's answer
        user_answer = st.text_area(
            "Provide your answer here:",
            value=st.session_state[f"user_answer_{current_index}"],
            key=f"user_answer_{current_index}",
            height=150
        )

        # Store the input from the user dynamically
        if st.button("Evaluate with AI"):
            if user_answer.strip() == "":
                st.error("Please provide an answer before evaluating.")
            else:
                with st.spinner("Evaluating your answer..."):
                    help_message = help_template.format(
                        question=current_question,
                        student_answer=user_answer,
                        ground_truth=current_answer
                    )
                    try:
                        help_response = help_llm(help_message)
                        st.write(help_response)
                    except Exception as e:
                        st.error(f"An error occurred during evaluation: {e}")

        if st.button("NEXT"):
            if user_answer.strip():
                st.session_state.count += 1
                st.rerun()
            else:
                st.error("Please provide an answer. If unsure, you can mention that you don't know.")

    else:
        st.markdown("## Congratulations, you've completed the session!")
        st.session_state.count = 0
        st.session_state.questions_generated = False

else:
    st.warning("Upload a document and specify the number of questions to proceed.")
